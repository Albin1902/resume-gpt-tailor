import streamlit as st
import re
import openai
from openai import OpenAI
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from io import BytesIO
from docx import Document

# Initialize OpenAI client
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# === Helper Functions ===
def extract_keywords(text):
    words = re.findall(r"\b[a-zA-Z]{4,}\b", text.lower())
    return sorted(set(words))

def extract_job_title_and_company(job_desc):
    job_title_match = re.search(r"(?i)(?:title|position|role)[:\-]?\s*(.+)", job_desc)
    job_title = job_title_match.group(1).strip() if job_title_match else ""
    if "your" in job_desc.lower() and "career" in job_desc.lower():
        lines = job_desc.splitlines()
        for i, line in enumerate(lines):
            if "your" in line.lower() and "career" in line.lower():
                job_title = lines[i].strip()
                break
    company_match = re.search(r"(?i)at ([A-Z][a-zA-Z0-9& ]+)", job_desc)
    company = company_match.group(1).strip() if company_match else ""
    if not company and "moneris" in job_desc.lower():
        company = "Moneris"
    return job_title, company

def extract_name_from_resume(resume_text):
    first_line = resume_text.strip().splitlines()[0]
    if len(first_line.split()) <= 5:
        return first_line.strip()
    return ""

def compute_ats_score(resume, job_desc):
    tfidf = TfidfVectorizer().fit_transform([resume, job_desc])
    score = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
    return round(score * 100, 2)

def clean_job_description(text):
    lines = text.splitlines()
    cutoff_keywords = ["why join us", "who we are", "what we offer", "equal opportunity", "requirements for success"]
    filtered = []
    for line in lines:
        if any(kw in line.lower() for kw in cutoff_keywords):
            break
        filtered.append(line)
    return "\n".join(filtered).strip()

def gpt_infer_role_tone(job_desc):
    prompt = f"""
Analyze the job description below and respond with:
1. Role title (e.g., Data Analyst)
2. Top 3 skills required
3. Ideal resume tone (e.g., formal, creative, data-driven, friendly)

Job Description:
{job_desc}
"""
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=300
    )
    return response.choices[0].message.content.strip()

def gpt_rewrite_resume(resume, job_desc, role_tone_summary, temperature):
    prompt = f"""
{role_tone_summary}

Now, rewrite the following resume to better match the job description. Emphasize the key skills and use the appropriate tone.

Resume:
{resume}

Job Description:
{job_desc}

Tailored Resume:
"""
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=1500
    )
    return response.choices[0].message.content.strip()

def gpt_generate_cover_letter(name, job_title, company, source, friend, temperature):
    referral_line = f" They were referred by {friend}." if source == "Referral" and friend else ""
    prompt = f"""
Write a personalized cover letter for {name} applying to the {job_title} role at {company}.
They found this job through {source}.{referral_line}
The tone should be confident, respectful, and focused on alignment with the role.

Cover Letter:
"""
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=700
    )
    return response.choices[0].message.content.strip()

def highlight_keywords(text, keywords):
    for kw in sorted(set(keywords), key=len, reverse=True):
        pattern = re.compile(rf"\b({re.escape(kw)})\b", re.IGNORECASE)
        text = pattern.sub(r"**\1**", text)
    return text

def text_to_docx(text, title):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# === Streamlit UI ===
st.title("GPT Resume & Cover Letter Tailor")
st.markdown("Paste your resume and job description to generate an ATS-optimized resume and cover letter.")

job_source = st.selectbox("How did you find the job?", ["LinkedIn", "Indeed", "Referral"])
friend_name = st.text_input("Friend's Name (if Referral)") if job_source == "Referral" else None

temp = st.slider("Select creativity (temperature)", 0.0, 1.0, 0.7, 0.1)

resume_input = st.text_area("Paste Your Resume", height=300)
job_input = st.text_area("Paste Job Description", height=300)

if st.button("Generate") and resume_input and job_input:
    user_name = extract_name_from_resume(resume_input)
    job_title, company = extract_job_title_and_company(job_input)
    ats_words = extract_keywords(job_input)
    cleaned_jd = clean_job_description(job_input)

    st.subheader("Detected Info")
    st.write(f"**Your Name:** {user_name or '❓ Not found'}")
    st.write(f"**Job Title:** {job_title or '❓ Not found'}")
    st.write(f"**Company Name:** {company or '❓ Not found'}")
    st.write(f"**ATS Keywords:** {', '.join(ats_words[:25])}")

    st.subheader("🔍 AI-Inferred Role and Resume Style")
    role_summary = gpt_infer_role_tone(job_input)
    st.info(role_summary)

    ats_score_before = compute_ats_score(resume_input, cleaned_jd)
    tailored_resume = gpt_rewrite_resume(resume_input, cleaned_jd, role_summary, temp)
    ats_score_after = compute_ats_score(tailored_resume, cleaned_jd)

    st.subheader("ATS Match Score")
    st.write(f"**Before Optimization:** {ats_score_before}%")
    st.write(f"**After Optimization:** {ats_score_after}%")

    st.subheader("Tailored Resume")
    st.text_area("Result", tailored_resume, height=400)
    st.download_button("Download Resume (.docx)", text_to_docx(tailored_resume, "Tailored Resume"), file_name="tailored_resume.docx")

    st.subheader("Cover Letter")
    cover_letter = gpt_generate_cover_letter(
        user_name or "Applicant", job_title or "the role", company or "the company", job_source, friend_name, temp
    )
    st.text_area("Result", cover_letter, height=300)
    st.download_button("Download Cover Letter (.docx)", text_to_docx(cover_letter, "Cover Letter"), file_name="cover_letter.docx")

    st.subheader("Matched Keywords in Resume")
    st.markdown(highlight_keywords(tailored_resume, ats_words), unsafe_allow_html=True)